#读取docx中的文本代码示例
import docx
import re

#获取文档
file = docx.Document("C:\\Users\\Administrator\\Desktop\\复习题.docx")
print("段落数:"+str(len(file.paragraphs))) #输出段落数

data_list = []
#输出每一段的内容
for para in file.paragraphs:
    #print(re.findall(r'\d+\、(.*?) .*', para.text))
    data_list.append(re.findall(r'\d+\、(.*?) .*', para.text))#未完成

while [] in data_list:
    data_list.remove([])
# print(data_list)

from selenium import webdriver
from selenium.webdriver.common.keys import Keys
# 引入Keys类包 发起键盘操作
import time

index = 0
j = index+1

for i in range(len(data_list)):
    print(j)
    j = j+1

    driver = webdriver.Chrome()
    # 访问百度
    driver.get('http://www.baidu.com')

    # 输入框输入内容
    driver.find_element_by_id('kw').send_keys(data_list[i+index-1])
    # 模拟回车操作 ,开始搜索
    driver.find_element_by_id('su').send_keys(Keys.ENTER)
    time.sleep(30)
    # 退出
    driver.quit()

